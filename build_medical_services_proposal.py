from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path("deliverables/MediRent_Flexible_Endoscopy_Service_Proposal.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)
LOGO_PATH = Path("/var/folders/sd/8mg7471n7n18b68sv4ffh_r40000gn/T/TemporaryItems/NSIRD_screencaptureui_05SYJr/Screenshot 2026-09-03 at 1.55.28 PM.png")

NAVY = "17324D"
TEAL = "147D82"
LIGHT = "EAF4F4"
PALE = "F5F8FA"
MID = "5A6773"
WHITE = "FFFFFF"
BLACK = "20262C"
GOLD = "C89B3C"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.82)
sec.left_margin = sec.right_margin = Inches(0.88)
sec.header_distance = sec.footer_distance = Inches(0.38)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.18

for name, size, before, after, color in [
    ("Title", 25, 0, 7, NAVY),
    ("Subtitle", 12, 0, 18, MID),
    ("Heading 1", 16, 16, 8, NAVY),
    ("Heading 2", 12.5, 11, 5, TEAL),
    ("Heading 3", 11.5, 8, 4, NAVY),
]:
    st = styles[name]
    st.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
    st._element.rPr.rFonts.set(qn("w:ascii"), st.font.name)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), st.font.name)
    st.font.size = Pt(size)
    st.font.bold = name != "Subtitle"
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=110, start=130, bottom=110, end=130):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = tc.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tc.append(tcMar)
    for tag, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        e = tcMar.find(qn("w:"+tag))
        if e is None:
            e = OxmlElement("w:"+tag); tcMar.append(e)
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")

def fix_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa"); tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i]); margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); trPr.append(rep)

def no_split(row):
    trPr = row._tr.get_or_add_trPr(); trPr.append(OxmlElement("w:cantSplit"))

def cell_text(cell, text, bold=False, color=BLACK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text); r.bold = bold; r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color)

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def add_callout(label, text, fill=LIGHT):
    t = doc.add_table(rows=1, cols=1)
    fix_table(t, [9360])
    c = t.cell(0,0); shade(c, fill); margins(c, 150, 180, 150, 180)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.12
    r = p.add_run(label + "  "); r.bold=True; r.font.color.rgb=RGBColor.from_string(TEAL)
    p.add_run(text)

# Running header and footer
hp = sec.header.paragraphs[0]
hp.text = "MEDIRENT  |  CLINICAL SERVICE PROPOSAL"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.name="Aptos"; r.font.size=Pt(8); r.bold=True; r.font.color.rgb=RGBColor.from_string(MID)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Confidential proposal  •  Prepared for review and approval")
fr.font.name="Aptos"; fr.font.size=Pt(8); fr.font.color.rgb=RGBColor.from_string(MID)

# Cover / opening
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
r = p.add_run()
r.add_picture(str(LOGO_PATH), width=Inches(1.85))

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Proposal for Flexible Endoscopy,\nLaser Urology and Men’s Health Services")
p = doc.add_paragraph(style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("On-site clinical lists, specialist support, equipment and sterile consumables")

t = doc.add_table(rows=3, cols=2)
fix_table(t, [4680,4680])
meta = [
    ("Prepared for", "[Hospital / Clinic Name]"),
    ("Prepared by", "MediRent / Dr. Muzaffar Chaudhary"),
    ("Status", "For review and approval"),
]
for i,(a,b) in enumerate(meta):
    shade(t.cell(i,0), PALE); cell_text(t.cell(i,0),a.upper(),True,MID,8.5)
    cell_text(t.cell(i,1),b, i==3, NAVY if i==3 else BLACK,9.5)

doc.add_paragraph()
add_callout("PROPOSAL SUMMARY", "MediRent proposes an initial complimentary trial followed by scheduled flexible cystoscopy, laser flexible ureteroscopy and men’s-health shockwave therapy services. Trained clinical support, specialist equipment, sterile instruments and consumables will be arranged as applicable to the agreed cases.")

doc.add_heading("1. Purpose", level=1)
doc.add_paragraph("This proposal is intended to establish a practical service arrangement through which the recipient hospital or clinic can offer diagnostic and therapeutic flexible endoscopy procedures, laser urology services and men’s-health shockwave therapy using visiting specialist support and rented equipment. The model is designed to begin with a limited trial, demonstrate workflow and clinical feasibility, and then move to scheduled paid services.")

doc.add_heading("2. Complimentary Trial", level=1)
doc.add_paragraph("MediRent proposes a trial list of three patients at no charge. Suitable trial cases may include:")
for item in ["Diagnostic flexible cystoscopy", "Ureteric stent removal", "Intravesical Botox injection"]: add_bullet(item)
doc.add_paragraph("A trained urologist will be available for the trial list. Patient selection, clinical suitability, consent and scheduling must be agreed in advance between the participating clinical teams.")
add_callout("IMPORTANT CLARIFICATION", "The source instruction states that the first three cases will be provided without charge. The parties should confirm in writing whether this waiver includes transport, all consumables and any hospital/facility charges before the trial date.", "FFF4D6")

doc.add_page_break()
doc.add_heading("3. Flexible Cystoscopy List", level=1)
doc.add_paragraph("A standard list is expected to accommodate approximately 8–12 patients within 3.5 hours, subject to case complexity, theatre efficiency, patient readiness and clinical requirements.")

table = doc.add_table(rows=1, cols=3)
fix_table(table,[4050,2250,3060]); set_repeat_header(table.rows[0])
for c,txt in zip(table.rows[0].cells,["Charge item","Rate","Basis / note"]): shade(c,NAVY); cell_text(c,txt,True,WHITE,9.5)
rows=[
    ("Transport charges","PKR 20,000","Per scheduled list"),
    ("Flexible scope and consumables","PKR 8,000","Per patient"),
]
for a,b,c in rows:
    cells=table.add_row().cells; no_split(table.rows[-1]);
    for cell in cells: shade(cell,WHITE)
    cell_text(cells[0],a); cell_text(cells[1],b,True,TEAL,9.5,WD_ALIGN_PARAGRAPH.CENTER); cell_text(cells[2],c)
doc.add_heading("Included for Each Scheduled List", level=2)
for item in [
    "Twelve individually packed sterile instruments and required consumables for the list.",
    "Flexible cystoscopy equipment and logistical transport, as agreed.",
    "A trained urologist for the proposed service period, subject to confirmed scheduling.",
]: add_bullet(item)

doc.add_heading("Minimum Booking", level=2)
doc.add_paragraph("A minimum of three booked patients is required for a list to be operationally sustainable. Final confirmation should be completed before equipment and personnel are dispatched.")

doc.add_heading("4. Laser Flexible Ureteroscopy", level=1)
table = doc.add_table(rows=1, cols=3)
fix_table(table,[4050,2250,3060]); set_repeat_header(table.rows[0])
for c,txt in zip(table.rows[0].cells,["Charge item","Rate","Basis / note"]): shade(c,NAVY); cell_text(c,txt,True,WHITE,9.5)
for a,b,c in [
    ("Transport charges","PKR 20,000","Per scheduled list / visit"),
    ("Laser and flexible ureteroscope rental","PKR 50,000","Per procedure"),
]:
    cells=table.add_row().cells; no_split(table.rows[-1]); cell_text(cells[0],a); cell_text(cells[1],b,True,TEAL,9.5,WD_ALIGN_PARAGRAPH.CENTER); cell_text(cells[2],c)

doc.add_heading("5. Men’s Health Shockwave Therapy", level=1)
doc.add_paragraph("MediRent can provide a shockwave therapy machine for rental and can also arrange treatment sessions for men’s health services, subject to clinical assessment and suitability.")
table = doc.add_table(rows=1, cols=3)
fix_table(table,[4050,2250,3060]); set_repeat_header(table.rows[0])
for c,txt in zip(table.rows[0].cells,["Service","Rate","Basis / note"]): shade(c,NAVY); cell_text(c,txt,True,WHITE,9.5)
for a,b,c in [
    ("Shockwave therapy procedure","PKR 15,000","Per treatment session"),
    ("Shockwave therapy machine rental","On quotation","Based on rental duration and location"),
]:
    cells=table.add_row().cells; no_split(table.rows[-1]); cell_text(cells[0],a); cell_text(cells[1],b,True,TEAL,9.5,WD_ALIGN_PARAGRAPH.CENTER); cell_text(cells[2],c)
doc.add_paragraph("Treatment plans, the number of sessions and patient eligibility will be determined by an appropriately qualified clinician. Pricing for machine-only rental will be provided after confirming the rental period, location, transport and support requirements.")

doc.add_heading("6. Responsibilities", level=1)
doc.add_heading("MediRent / Visiting Team", level=2)
for item in [
    "Provide the agreed scope, equipment, sterile instruments and consumables.",
    "Arrange transport and deliver the equipment for the confirmed list.",
    "Provide appropriately trained urology support as agreed for the trial and subsequent lists.",
    "Coordinate equipment setup, handling and post-list reconciliation.",
]: add_bullet(item)
doc.add_heading("Recipient Hospital / Clinic", level=2)
for item in [
    "Identify, assess and schedule clinically suitable patients.",
    "Provide a credentialed doctor/urologist and appropriate clinical staff for future lists.",
    "Provide the clinical area, standard monitoring, medicines, emergency readiness, sterilization support and waste disposal required by hospital policy.",
    "Manage patient records, informed consent, investigations, clinical governance and follow-up care.",
    "Confirm the final list and patient count before dispatch of personnel and equipment.",
]: add_bullet(item)

doc.add_heading("7. Proposed Implementation", level=1)
steps=[
    ("Clinical alignment", "Confirm eligible procedures, responsible clinicians, facility requirements and governance approvals."),
    ("Trial list", "Schedule up to three suitable patients and conduct the complimentary demonstration list."),
    ("Review", "Evaluate workflow, safety, patient throughput, equipment performance and staff requirements."),
    ("Routine scheduling", "Agree list frequency, booking cut-off, payment process and cancellation terms."),
]
for i,(title,body) in enumerate(steps,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.08); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(f"{i}. {title}: "); r.bold=True; r.font.color.rgb=RGBColor.from_string(TEAL); p.add_run(body)

doc.add_heading("8. Commercial and Clinical Conditions", level=1)
for item in [
    "All amounts are stated in Pakistani rupees (PKR) and should be confirmed before signing.",
    "The final quotation will depend on confirmed procedure type, patient count and any additional consumables or facility requirements.",
    "Patient selection and all clinical decisions remain the responsibility of appropriately credentialed clinicians.",
    "The parties should document payment timing, taxes, cancellation/rescheduling terms, equipment damage or loss, infection-control responsibilities and indemnity before routine service begins.",
    "This proposal is a commercial service outline and does not replace a clinical protocol, credentialing process or formal service agreement.",
]: add_bullet(item)

doc.add_heading("9. Acceptance in Principle", level=1)
doc.add_paragraph("If the above framework is acceptable, the parties may proceed to confirm the trial date, patient list, participating clinicians and facility readiness. Final commercial and clinical terms should be recorded in a signed service agreement or approved purchase order.")

sig = doc.add_table(rows=4, cols=2)
fix_table(sig,[4680,4680])
labels=[("For the Recipient Hospital / Clinic","For MediRent / Service Provider"),("Name: __________________________","Name: __________________________"),("Designation: ____________________","Designation: ____________________"),("Signature & Date: ________________","Signature & Date: ________________")]
for i,(a,b) in enumerate(labels):
    if i==0: shade(sig.cell(i,0),LIGHT); shade(sig.cell(i,1),LIGHT)
    cell_text(sig.cell(i,0),a,i==0,NAVY if i==0 else BLACK,9.5)
    cell_text(sig.cell(i,1),b,i==0,NAVY if i==0 else BLACK,9.5)

# Keep rows together and normalize table paragraph fonts
for table in doc.tables:
    for row in table.rows:
        no_split(row)

doc.core_properties.title = "Proposal for Flexible Endoscopy, Laser Urology and Men’s Health Services"
doc.core_properties.subject = "Clinical service and equipment rental proposal"
doc.core_properties.author = "MediRent"
doc.core_properties.keywords = "MediRent, flexible cystoscopy, flexible ureteroscopy, laser urology"
doc.save(OUT)
print(OUT.resolve())
